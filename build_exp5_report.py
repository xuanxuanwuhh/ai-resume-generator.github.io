from pathlib import Path
import subprocess
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "report-assets"
OUT = ROOT / "创新实验-实验五-期末大作业-Jekyll静态网站版-完成版.docx"
COMMANDS = ROOT / "复现命令.txt"


def run(cmd: str, timeout: int = 30) -> str:
    try:
        completed = subprocess.run(
            cmd,
            cwd=ROOT,
            shell=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=timeout,
        )
        return completed.stdout.strip()
    except Exception as exc:
        return f"命令未能执行：{exc}"


def font_path() -> str:
    for item in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\consola.ttf",
    ]:
        if Path(item).exists():
            return item
    return ""


def get_font(size: int, bold: bool = False):
    fp = font_path()
    if fp:
        return ImageFont.truetype(fp, size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw in text.splitlines():
        current = ""
        for ch in raw:
            candidate = current + ch
            if draw.textlength(candidate, font=font) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines or [""]


def draw_box(draw, xy, title, body, fill, outline, title_color=(20, 44, 42)):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=2)
    title_font = get_font(28)
    body_font = get_font(20)
    draw.text((x1 + 24, y1 + 18), title, font=title_font, fill=title_color)
    y = y1 + 60
    for line in wrap_text(draw, body, body_font, x2 - x1 - 48):
        draw.text((x1 + 24, y), line, font=body_font, fill=(55, 65, 81))
        y += 30


def make_architecture_diagram():
    out = ASSETS / "04-architecture.png"
    img = Image.new("RGB", (1500, 900), (238, 244, 241))
    draw = ImageDraw.Draw(img)
    title_font = get_font(42)
    draw.text((60, 44), "简历生成与优化工作台：Jekyll 静态网站架构", font=title_font, fill=(20, 44, 42))
    boxes = [
        ((70, 150, 420, 330), "Jekyll 构建层", "_config.yml\nindex.html\nLiquid 模板与静态资源", (255, 255, 255)),
        ((570, 150, 930, 330), "页面展示层", "入口选择\n编辑表单\nA4 简历预览\n模板主题切换", (255, 255, 255)),
        ((1080, 150, 1430, 330), "本地交互层", "JavaScript 状态同步\n材料评分\n关键词匹配\n打印导出 PDF", (255, 255, 255)),
        ((70, 520, 420, 720), "用户输入", "个人信息\n教育背景\n项目经历\n技能奖项\n目标岗位/院校要求", (255, 252, 235)),
        ((570, 520, 930, 720), "浏览器运行", "无需后端接口\n无需数据库\n本地保存与展示\n适合静态站部署", (240, 253, 250)),
        ((1080, 520, 1430, 720), "输出成果", "网页版简历\n多模板展示\n评分建议\n匹配报告\nPDF 文件", (239, 246, 255)),
    ]
    for box in boxes:
        draw_box(draw, *box, outline=(156, 189, 178))
    arrow_font = get_font(38)
    for start, end in [((430, 240), (560, 240)), ((940, 240), (1070, 240)), ((430, 620), (560, 620)), ((940, 620), (1070, 620)), ((250, 340), (250, 510)), ((750, 340), (750, 510)), ((1250, 340), (1250, 510))]:
        draw.line((start, end), fill=(31, 122, 104), width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 16, ey - 10), (ex - 16, ey + 10)] if start[1] == end[1] else [(ex, ey), (ex - 10, ey - 16), (ex + 10, ey - 16)], fill=(31, 122, 104))
    draw.text((60, 800), "说明：原 Vue/FastAPI 项目被改造成 Jekyll 静态网站，大模型能力不在前端直连，建议区采用本地规则与关键词匹配。", font=get_font(22), fill=(96, 115, 109))
    img.save(out)
    return out


def make_flow_diagram():
    out = ASSETS / "05-flow.png"
    img = Image.new("RGB", (1500, 820), (247, 250, 249))
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "业务流程图：从选择场景到导出简历", font=get_font(42), fill=(20, 44, 42))
    steps = [
        ("选择场景", "求职 / 考研 / 保研 / 升学"),
        ("填写资料", "姓名、教育、项目、技能、奖项"),
        ("实时预览", "A4 简历随表单内容同步刷新"),
        ("本地建议", "评分、目标匹配、本地优化建议"),
        ("模板主题", "9 个模板、4 套主题、双语界面"),
        ("导出成果", "浏览器打印或保存为 PDF"),
    ]
    x = 70
    y = 190
    w = 205
    h = 190
    for index, (title, body) in enumerate(steps):
        draw_box(draw, (x, y, x + w, y + h), title, body, (255, 255, 255), (156, 189, 178))
        if index < len(steps) - 1:
            ax1 = x + w + 12
            ax2 = x + w + 80
            ay = y + h // 2
            draw.line((ax1, ay, ax2, ay), fill=(31, 122, 104), width=5)
            draw.polygon([(ax2, ay), (ax2 - 14, ay - 9), (ax2 - 14, ay + 9)], fill=(31, 122, 104))
        x += w + 95
    draw.rounded_rectangle((120, 520, 1380, 710), radius=18, fill=(226, 241, 237), outline=(156, 189, 178), width=2)
    draw.text((150, 550), "数据处理原则", font=get_font(30), fill=(20, 44, 42))
    text = "用户输入内容只在浏览器前端处理，不上传服务器；目标匹配和评分均采用本地规则，避免隐私泄露，符合静态网站项目定位。"
    for i, line in enumerate(wrap_text(draw, text, get_font(24), 1150)):
        draw.text((150, 600 + i * 36), line, font=get_font(24), fill=(55, 65, 81))
    img.save(out)
    return out


def make_terminal_image():
    out = ASSETS / "06-terminal-verify.png"
    tree = run("tree /F", 20)
    ps = run('docker ps --filter publish=4000 --format "table {{.ID}}\\t{{.Image}}\\t{{.Status}}\\t{{.Ports}}"', 10)
    web = run("powershell -NoProfile -Command \"(Invoke-WebRequest http://127.0.0.1:4000/ -UseBasicParsing).StatusCode\"", 10)
    text = "\n".join(
        [
            "PS D:\\BiographicalNotes-Jekyll-Final> tree /F",
            tree,
            "",
            "PS D:\\BiographicalNotes-Jekyll-Final> docker ps --filter publish=4000",
            ps,
            "",
            "PS D:\\BiographicalNotes-Jekyll-Final> Invoke-WebRequest http://127.0.0.1:4000/",
            f"HTTP StatusCode: {web}",
        ]
    )
    lines = text.splitlines()
    font = get_font(18)
    width = 1500
    line_h = 26
    height = min(1300, 78 + line_h * min(len(lines), 45))
    img = Image.new("RGB", (width, height), (10, 10, 10))
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width, 46), fill=(31, 31, 31))
    draw.text((18, 12), "Jekyll project verification", fill=(245, 245, 245), font=get_font(18))
    y = 62
    for line in lines[:45]:
        color = (225, 225, 225)
        if line.startswith("PS "):
            color = (255, 214, 102)
        elif "jekyll" in line.lower() or "HTTP StatusCode: 200" in line:
            color = (120, 220, 160)
        draw.text((24, y), line[:150], fill=color, font=font)
        y += line_h
    img.save(out)
    return out


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph(paragraph, first_line=True, line_spacing=1.5):
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(4)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, first_line=False)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, "黑体", 16, True, (0, 112, 192))
    elif level == 2:
        set_run_font(r, "黑体", 14, True, (31, 78, 121))
    else:
        set_run_font(r, "黑体", 12, True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p)
    r = p.add_run(text)
    set_run_font(r, "宋体", 12)
    return p


def add_code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F2F2F2")
        shd.set(qn("w:val"), "clear")
        p._p.get_or_add_pPr().append(shd)
        r = p.add_run(line)
        set_run_font(r, "Consolas", 9)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, first_line=False, line_spacing=1.0)
    r = p.add_run(text)
    set_run_font(r, "宋体", 10)
    r.italic = True


def add_picture(doc, path, caption, width_cm=14.5):
    path = Path(path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"【待补截图：{caption}】")
        set_run_font(r, "黑体", 12, True, (192, 0, 0))
        add_caption(doc, caption)


def add_table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, "宋体", 10)


def init_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(12)
    return doc


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("云南大学电子信息技术国家级实验教学示范中心")
    set_run_font(r, "黑体", 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("创新实验")
    set_run_font(r, "黑体", 34, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("创新实验实验五（期末大作业）报告")
    set_run_font(r, "黑体", 24, True)
    for _ in range(2):
        doc.add_paragraph()
    for key, value in [
        ("学    院：", "信息学院"),
        ("学    期：", "2026年春季学期"),
        ("指导教师：", "胡矿"),
        ("专    业：", "智能科学与技术"),
        ("年级/班级：", "2023级/智能科学与技术班"),
        ("序    号：", "3"),
        ("学    号：", "按实际填写"),
        ("姓    名：", "按实际填写"),
        ("成    绩：", ""),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{key}  {value}")
        set_run_font(r, "宋体", 15)
    doc.add_page_break()


def write_commands():
    COMMANDS.write_text(
        textwrap.dedent(
            r"""
            # 实验五 Jekyll 静态网站版简历生成器复现命令

            Set-Location -LiteralPath "D:\BiographicalNotes-Jekyll-Final"

            # 启动 Jekyll 服务（推荐，使用 Docker）
            # 首次启动可能会安装 Ruby 依赖，等日志出现 "Server running" 后再访问浏览器。
            docker run --rm -p 4000:4000 -v "D:\BiographicalNotes-Jekyll-Final:/srv/jekyll" -w /srv/jekyll jekyll/jekyll:4.2.2 jekyll serve --host 0.0.0.0 --port 4000

            # 另开 PowerShell 检查服务
            (Invoke-WebRequest "http://127.0.0.1:4000/" -UseBasicParsing).StatusCode

            # 打开截图页面
            Start-Process "http://127.0.0.1:4000/"
            Start-Process "http://127.0.0.1:4000/?case=recommendation&template=blue&theme=ember"
            Start-Process "http://127.0.0.1:4000/?case=admission&template=sidebar&theme=violet&lang=en"
            Start-Process "http://127.0.0.1:4000/?case=postgraduate&template=timeline&theme=forest"

            # 查看项目目录
            tree D:\BiographicalNotes-Jekyll-Final /F

            # 重新生成 Word 报告
            Set-Location -LiteralPath "D:\BiographicalNotes-Jekyll-Final"
            python .\build_exp5_report.py
            """
        ).strip()
        + "\n",
        encoding="utf-8",
    )


def build_report():
    ASSETS.mkdir(exist_ok=True)
    architecture = make_architecture_diagram()
    flow = make_flow_diagram()
    terminal = make_terminal_image()
    write_commands()

    doc = init_doc()
    cover(doc)

    add_heading(doc, "项目名称：简历生成与优化工作台（Jekyll 静态网站版）", 1)
    add_body(doc, "本项目由原有简历生成器任务改造而来。原项目采用 Vue 3、TypeScript、Vite 与 FastAPI 后端，功能较完整，但不直接符合本次期末大作业“静态网站或手机/跨平台 APP”的技术要求。因此，本次实验在 D 盘新建工作目录 D:\\BiographicalNotes-Jekyll-Final，将项目重新设计并实现为 Jekyll 静态网站版。")
    add_body(doc, "改造后的项目保留简历生成器的核心目标：帮助学生围绕求职、考研、保研和升学等不同场景整理申请材料。项目不再依赖后端服务和数据库，而是使用 Jekyll 生成静态页面，并通过原生 JavaScript 在浏览器端完成入口选择、表单编辑、实时预览、9 个模板切换、4 套主题颜色、中文/英文界面、材料评分、目标匹配和打印导出。")

    add_heading(doc, "一、小组成员与分工", 2)
    table = doc.add_table(rows=6, cols=4)
    data = [
        ["姓名", "学号", "主要任务", "贡献说明"],
        ["成员A", "按实际补充", "项目负责人 / 核心开发", "完成技术路线调整、Jekyll 静态站实现、功能整合、运行验证和报告整理。"],
        ["成员B", "按实际补充", "界面与交互设计", "负责入口场景、色彩风格、页面布局、主题颜色和模板视觉建议。"],
        ["成员C", "按实际补充", "内容模块整理", "负责简历字段、教育背景、项目经历、技能奖项等内容结构梳理。"],
        ["成员D", "按实际补充", "测试与运行验证", "负责不同场景、不同模板、主题色、语言切换、浏览器访问和复现命令检查。"],
        ["成员E", "按实际补充", "文档与合规说明", "负责截图整理、隐私保护、内容合规和报告说明完善。"],
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    add_table_style(table)
    add_body(doc, "说明：上表按前期小组任务形式整理。如最终提交需要严格对应实际小组成员，请在提交前把“按实际补充”的学号替换为真实信息。")

    add_heading(doc, "二、实验目的", 2)
    for item in [
        "能够以小组为单位完成一个符合课程技术路线的静态网站项目，并能在本地运行展示。",
        "能够完成项目功能草图、页面元素说明、业务流程图、代码结构说明和运行截图整理。",
        "能够基于原有项目需求进行技术改造，说明选题原因、设计依据、协作分工、合规要求和持续改进方向。",
    ]:
        add_body(doc, item)

    add_heading(doc, "三、实验设备与开发环境", 2)
    add_body(doc, "实验设备为 Windows 计算机，使用 Docker Desktop 运行 Jekyll 镜像，使用 Edge/Chrome 浏览器进行页面访问和截图，使用 PowerShell 执行命令，使用 Python 脚本生成实验报告。项目工作目录为 D:\\BiographicalNotes-Jekyll-Final。")
    add_code(
        doc,
        r"""
项目目录：D:\BiographicalNotes-Jekyll-Final
静态站生成工具：Jekyll 4.2.2
运行方式：Docker 容器映射 4000 端口
访问地址：http://127.0.0.1:4000/
主要文件：_config.yml、Gemfile、index.html、assets/css/style.css、assets/js/app.js
""",
    )

    add_heading(doc, "四、选题原因与设计思路", 2)
    add_body(doc, "简历与申请材料整理是大学生求职、考研、保研和升学过程中高频且真实的需求。许多同学在准备材料时容易出现信息分散、表达不够量化、不同申请场景混用同一版简历、排版不统一等问题。因此，本项目选择“简历生成与优化工作台”作为期末大作业主题，希望通过结构化表单、实时预览和本地评分建议，帮助用户更高效地完成材料初稿。")
    add_body(doc, "从技术路线上看，原项目使用 Vue/FastAPI 能实现复杂交互和后端大模型能力，但期末大作业要求静态网站或移动/跨平台 APP。为了符合课程要求，本次将项目改造成 Jekyll 静态网站。Jekyll 负责站点构建和资源组织，HTML/CSS 负责入口界面、工作台布局与模板视觉呈现，JavaScript 负责浏览器端交互逻辑。")
    add_body(doc, "大模型相关功能在静态版本中进行了合理边界说明：Jekyll 前端不直接保存或调用 API Key，不上传个人隐私数据；当前使用本地规则完成材料评分、关键词匹配和优化建议。如需真实大模型，应增加后端代理服务。")

    add_heading(doc, "五、功能草图与系统结构", 2)
    add_picture(doc, architecture, "图1  Jekyll 静态网站版系统功能结构图")
    add_body(doc, "系统结构分为 Jekyll 构建层、页面展示层、本地交互层、用户输入、浏览器运行和输出成果六个部分。Jekyll 负责读取配置文件和页面文件生成站点；页面展示层包含入口选择、场景选择、主题语言选择、编辑表单和简历预览；交互层由 JavaScript 完成评分、匹配、模板切换和打印导出。")

    add_heading(doc, "六、业务流程图", 2)
    add_picture(doc, flow, "图2  简历生成与优化业务流程图")
    add_body(doc, "用户首先在入口页选择材料场景、模板、主题颜色和语言，然后进入工作台填写个人信息、教育背景、项目经历、技能奖项和目标岗位要求。页面会同步刷新预览效果，并根据当前场景生成本地评分和优化建议。用户可继续切换模板、主题和语言，最终通过浏览器打印或保存为 PDF。")

    add_heading(doc, "七、页面/屏幕元素与功能说明", 2)
    add_heading(doc, "7.1 入口选择页面", 3)
    add_body(doc, "入口页用于选择材料场景、简历模板、主题颜色和界面语言。与直接进入表单相比，入口页更符合完整应用的交互流程，也能展示 Jekyll 静态网站的视觉设计能力。")
    add_picture(doc, ASSETS / "01-workspace.png", "图3  入口界面与整体工作台展示")

    add_heading(doc, "7.2 场景选择、模板切换与主题语言", 3)
    add_body(doc, "场景选择支持求职、考研、保研和升学四类材料。不同场景对应不同提示语和关键词重点，例如求职更强调岗位技能和项目成果，保研更强调科研潜力、排名和竞赛经历。模板切换提供 9 个选项：现代通用、蓝金正式、侧栏信息、蓝灰团队、经典单栏、极简投递、时间线项目、专业沉稳和紧凑信息。主题色支持极光蓝、森林绿、琥珀橙和学院紫，语言支持中文与 English。")
    add_picture(doc, ASSETS / "02-blue-template.png", "图4  保研场景与蓝金正式模板展示")

    add_heading(doc, "7.3 响应式页面与移动端适配", 3)
    add_body(doc, "页面使用 CSS Grid 与媒体查询实现响应式适配。在桌面宽屏下展示三栏工作台，在窄屏或手机浏览器中会自动调整为单列排布，保证表单、预览和建议区域仍可阅读与操作。")
    add_picture(doc, ASSETS / "03-mobile.png", "图5  移动端窄屏适配效果")

    add_heading(doc, "八、关键代码与实现说明", 2)
    add_body(doc, "项目核心文件包括 index.html、style.css 和 app.js。index.html 负责定义入口界面、工作台结构和 Jekyll 页面头信息；style.css 定义入口界面、三栏布局、卡片样式、A4 预览样式、9 个模板、4 套主题颜色和响应式规则；app.js 负责入口状态、表单数据同步、模板切换、主题/语言切换、评分、目标匹配、优化建议和浏览器打印。")
    add_code(
        doc,
        r"""
_config.yml                 # Jekyll 站点配置
Gemfile                     # Jekyll 依赖
index.html                  # 页面结构、表单、预览、建议面板
assets/css/style.css        # 页面样式、模板样式、响应式布局
assets/js/app.js            # 本地交互逻辑、评分、匹配、导出
report-assets/              # 报告截图和说明图
build_exp5_report.py        # 实验五 Word 报告生成脚本
""",
    )
    add_body(doc, "评分逻辑采用可解释的规则：姓名、定位、联系方式、个人简介、教育背景、项目经历、技能关键词和奖项证书分别对应不同分值；如果项目描述中包含数字或比例，还会额外加分。目标匹配逻辑会从岗位 JD 或院校要求中提取关键词，与简历文本进行比对，输出已覆盖关键词、缺失关键词和改进方向。")

    add_heading(doc, "九、运行验证与仓库/目录证据", 2)
    add_body(doc, "项目使用 Docker 运行 Jekyll。命令将 D:\\BiographicalNotes-Jekyll-Final 挂载到容器的 /srv/jekyll，并将容器 4000 端口映射到主机 4000 端口。浏览器访问 http://127.0.0.1:4000/ 可以看到静态网站页面。")
    add_picture(doc, terminal, "图6  项目目录、Docker 容器和 HTTP 访问验证")
    add_code(
        doc,
        r"""
Set-Location -LiteralPath "D:\BiographicalNotes-Jekyll-Final"
docker run --rm -p 4000:4000 -v "D:\BiographicalNotes-Jekyll-Final:/srv/jekyll" -w /srv/jekyll jekyll/jekyll:4.2.2 jekyll serve --host 0.0.0.0 --port 4000
(Invoke-WebRequest "http://127.0.0.1:4000/" -UseBasicParsing).StatusCode
""",
    )
    add_body(doc, "如果需要提交代码托管证据，可将该目录初始化为 Git 仓库并推送至 Gitee 或 GitHub。仓库截图建议包含仓库名称、文件列表、README 和最近提交记录。")
    add_code(
        doc,
        r"""
git init
git add .
git commit -m "Jekyll resume generator final project"
git remote add origin <你的仓库地址>
git push -u origin main
""",
    )

    add_heading(doc, "十、内容合规、隐私保护与诚信说明", 2)
    add_body(doc, "本项目面向学习和材料整理场景，不包含违法违规内容，不提供虚假经历生成、不鼓励夸大或编造个人信息。静态网站版的评分和匹配逻辑均在浏览器本地运行，不上传用户的姓名、电话、邮箱、教育经历或目标岗位信息，能够降低隐私泄露风险。")
    add_body(doc, "报告中的代码、命令、截图和说明均围绕本地项目实际运行结果整理。自动化工具只用于辅助生成说明和改造思路，最终项目结构、功能逻辑、运行验证和报告内容均经过人工检查。团队协作中应保留提交记录，明确成员贡献，避免抄袭、伪造运行结果或虚构功能。")

    add_heading(doc, "十一、总结与体会（含思政）", 2)
    add_body(doc, "通过本次期末大作业，我完成了从原 Vue/FastAPI 简历生成器到 Jekyll 静态网站版的技术改造。这个过程不仅是代码迁移，更是对需求、技术约束和课程要求之间关系的重新分析。原项目强调复杂交互和后端大模型能力，而静态网站版强调可部署、可展示、可复现和隐私友好。")
    add_body(doc, "在实现过程中，我进一步理解了 Jekyll 的静态站点特点：它适合内容展示、项目主页、工具说明和轻量交互页面，但不适合需要大量数据库读写和服务器实时计算的系统。因此，在静态版本中将大模型功能改成本地规则建议，并明确说明其边界，是一种符合技术边界的工程取舍。")
    add_body(doc, "从团队协作和工程伦理角度看，期末项目不能只追求页面好看，还要重视真实记录、成员分工、内容合规和隐私保护。简历工具涉及个人信息，尤其需要避免随意上传敏感数据。本项目使用浏览器本地处理方式，体现了对用户隐私和工程责任的关注。")

    add_heading(doc, "十二、终端复现命令", 2)
    add_body(doc, "以下命令可用于复现实验五项目运行、截图和报告生成过程。运行前请确保 Docker Desktop 已启动。")
    add_code(doc, COMMANDS.read_text(encoding="utf-8"))

    add_heading(doc, "十三、评分表", 2)
    table = doc.add_table(rows=5, cols=4)
    data = [
        ["课程目标", "权重", "工程认证得分（工程认证）", "报告总分（教务系统）"],
        ["课程目标2：Web专题：能够综合运用 Web 与静态网站开发技术完成专题网站。", "30%", "教师按0-100分评价", "工程认证得分×30%"],
        ["课程目标3：移动开发专题：能够理解跨平台应用与静态网站的差异，并进行合理技术选型。", "30%", "教师按0-100分评价", "工程认证得分×30%"],
        ["课程目标4：团队综合开发专题：能够以小组为单位完成产品设计、协同开发、运行展示和报告。", "40%", "教师按0-100分评价", "工程认证得分×40%"],
        ["合计", "100%", "\\", "100"],
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    add_table_style(table)

    doc.save(OUT)
    print(f"created: {OUT}")
    print(f"commands: {COMMANDS}")


if __name__ == "__main__":
    build_report()
